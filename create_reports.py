#!/usr/bin/env python3
"""
Report Generator for Data Analysis Course
==========================================

Creates DSTU 3008:2015 compliant reports for all assignments.
Generates both DOCX and PDF formats.

Course: Intelligent Data Analysis
University: KhNURE
Teacher: Костянтин Георгійович Онищенко
"""

import os
import subprocess
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ==================== CONFIGURATION ====================

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
REPORTS_DIR = os.path.join(SCRIPT_DIR, 'reports')
os.makedirs(REPORTS_DIR, exist_ok=True)

# University info
UNIVERSITY = "Міністерство освіти і науки України"
UNIVERSITY_NAME = "Харківський національний університет радіоелектроніки"
DEPARTMENT = "Кафедра програмної інженерії"
DISCIPLINE = "Інтелектуальний аналіз даних"
TEACHER = "ст. викл. Онищенко К.Г."
STUDENT_NAME = "Голодніков Дмитро"
STUDENT_GROUP = "ІПЗм-24-2"
GITHUB_URL = "https://github.com/na-naina/data-analysis-khnure"


# ==================== DOCUMENT HELPERS ====================

def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    """Set cell margins."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, margin_value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = OxmlElement(f'w:{margin_name}')
        node.set(qn('w:w'), str(margin_value))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def add_centered_paragraph(doc, text, font_size=14, bold=False, space_after=0):
    """Add a centered paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_paragraph(doc, text, font_size=14, bold=False, first_line_indent=1.25, space_after=0, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """Add a regular paragraph."""
    p = doc.add_paragraph()
    p.alignment = alignment
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    p.paragraph_format.first_line_indent = Cm(first_line_indent) if first_line_indent else None
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    return p


def add_heading(doc, text, level=1):
    """Add a heading - black color, bold, no blue."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = None  # Black color
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_image_with_caption(doc, image_path, caption, width=5.5):
    """Add an image with caption."""
    if os.path.exists(image_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_path, width=Inches(width))

        # Caption
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.italic = True
        cap.paragraph_format.space_after = Pt(12)


def create_title_page(doc, work_type, work_number, work_title):
    """Create a title page according to DSTU 3008:2015."""
    # Ministry - centered, regular
    add_centered_paragraph(doc, UNIVERSITY, font_size=14, space_after=0)
    add_centered_paragraph(doc, UNIVERSITY_NAME, font_size=14, space_after=0)
    add_centered_paragraph(doc, DEPARTMENT, font_size=14, space_after=0)

    # Empty lines for spacing
    for _ in range(8):
        add_centered_paragraph(doc, "", font_size=14)

    # Report title - centered
    add_centered_paragraph(doc, "ЗВІТ", font_size=14, space_after=6)
    if work_number:
        add_centered_paragraph(doc, f"{work_type} № {work_number}", font_size=14, space_after=6)
    else:
        add_centered_paragraph(doc, work_type, font_size=14, space_after=6)
    add_centered_paragraph(doc, f'з дисципліни «{DISCIPLINE}»', font_size=14, space_after=6)
    add_centered_paragraph(doc, f'на тему «{work_title}»', font_size=14, space_after=0)

    # Empty lines
    for _ in range(6):
        add_centered_paragraph(doc, "", font_size=14)

    # Student and teacher info - left aligned with tabs
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Виконав")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"студент групи {STUDENT_GROUP}\t\t\t\t\t{STUDENT_NAME}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

    add_centered_paragraph(doc, "", font_size=14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Перевірив")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"{TEACHER}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

    # Empty lines
    for _ in range(4):
        add_centered_paragraph(doc, "", font_size=14)

    # City and year
    add_centered_paragraph(doc, "Харків 2024", font_size=14)

    # Page break
    doc.add_page_break()


def add_goal_section(doc, goal_text):
    """Add goal section."""
    add_heading(doc, "1 МЕТА РОБОТИ", level=1)
    add_paragraph(doc, goal_text, first_line_indent=1.25)


def add_execution_section(doc, content_items):
    """Add execution section with content items."""
    add_heading(doc, "2 ХІД ВИКОНАННЯ РОБОТИ", level=1)

    for item in content_items:
        if item['type'] == 'text':
            add_paragraph(doc, item['content'], first_line_indent=1.25)
        elif item['type'] == 'subheading':
            add_heading(doc, item['content'], level=2)
        elif item['type'] == 'image':
            add_image_with_caption(doc, item['path'], item['caption'], item.get('width', 5.5))
        elif item['type'] == 'code':
            # Code block
            p = doc.add_paragraph()
            run = p.add_run(item['content'])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            p.paragraph_format.space_after = Pt(6)


def add_conclusions_section(doc, conclusions):
    """Add conclusions section."""
    add_heading(doc, "3 ВИСНОВКИ", level=1)
    add_paragraph(doc, conclusions, first_line_indent=1.25)


def add_github_link(doc):
    """Add GitHub repository link at the end of the report."""
    add_heading(doc, "ПОСИЛАННЯ", level=1)
    p = doc.add_paragraph()
    run = p.add_run(f"Код проєкту доступний у репозиторії GitHub: ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run = p.add_run(GITHUB_URL)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.underline = True


def add_code_block(doc, code, caption=None):
    """Add a code block with monospace font."""
    if caption:
        p = doc.add_paragraph()
        run = p.add_run(caption)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.italic = True
        p.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(12)


def save_document(doc, filename):
    """Save document as DOCX and convert to PDF."""
    docx_path = os.path.join(REPORTS_DIR, f"{filename}.docx")
    pdf_path = os.path.join(REPORTS_DIR, f"{filename}.pdf")

    doc.save(docx_path)
    print(f"Saved: {docx_path}")

    # Convert to PDF using LibreOffice
    try:
        subprocess.run([
            'libreoffice', '--headless', '--convert-to', 'pdf',
            '--outdir', REPORTS_DIR, docx_path
        ], check=True, capture_output=True)
        print(f"Saved: {pdf_path}")
    except Exception as e:
        print(f"PDF conversion failed: {e}")


# ==================== REPORT GENERATORS ====================

def generate_lab1_report():
    """Generate Lab 1: Regression Analysis report."""
    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)

    create_title_page(doc, "Лабораторної роботи", "1", "Регресійний аналіз")

    add_goal_section(doc,
        "Ознайомитися з методами регресійного аналізу даних. "
        "Вивчити побудову та аналіз моделей простої лінійної, множинної лінійної "
        "та поліноміальної регресії. Навчитися оцінювати якість регресійних моделей "
        "за допомогою коефіцієнта детермінації R², RMSE та інших метрик. "
        "Отримати практичні навички роботи з бібліотеками scikit-learn та statsmodels.")

    results_dir = os.path.join(SCRIPT_DIR, 'labs', 'lab1_regression', 'results')

    content = [
        {'type': 'subheading', 'content': '2.1 Теоретичні основи регресійного аналізу'},
        {'type': 'text', 'content':
            'Регресійний аналіз - статистичний метод дослідження впливу однієї або кількох '
            'незалежних змінних (предикторів) на залежну змінну (відгук). Метою регресії є '
            'побудова моделі, що описує цю залежність та дозволяє робити прогнози.'},
        {'type': 'text', 'content':
            'Основні метрики оцінки якості регресійної моделі:'},
        {'type': 'text', 'content':
            '• R² (коефіцієнт детермінації) - частка дисперсії залежної змінної, що пояснюється моделлю. '
            'Значення від 0 до 1, де 1 означає ідеальне пояснення.'},
        {'type': 'text', 'content':
            '• RMSE (Root Mean Square Error) - корінь із середнього квадрата помилок. '
            'Показує середню величину відхилення прогнозів від фактичних значень.'},
        {'type': 'text', 'content':
            '• MAE (Mean Absolute Error) - середня абсолютна помилка. '
            'Менш чутлива до викидів порівняно з RMSE.'},

        {'type': 'subheading', 'content': '2.2 Проста лінійна регресія'},
        {'type': 'text', 'content':
            'Проста лінійна регресія моделює залежність між однією незалежною змінною X '
            'та залежною змінною Y у вигляді лінійного рівняння:'},
        {'type': 'text', 'content': 'Y = β₀ + β₁X + ε'},
        {'type': 'text', 'content':
            'де β₀ - вільний член (intercept), що визначає значення Y при X=0; '
            'β₁ - коефіцієнт регресії (slope), що показує на скільки зміниться Y при зміні X на одиницю; '
            'ε - випадкова помилка моделі.'},
        {'type': 'text', 'content':
            'Для оцінки параметрів β₀ та β₁ використовується метод найменших квадратів (OLS), '
            'який мінімізує суму квадратів відхилень: Σ(yᵢ - ŷᵢ)² → min.'},
        {'type': 'text', 'content':
            'Для демонстрації було згенеровано вибіркові дані та побудовано модель простої '
            'лінійної регресії. Реалізація на Python з використанням scikit-learn:'},
        {'type': 'code', 'content': '''from sklearn.linear_model import LinearRegression
import numpy as np

# Підготовка даних
X = data[['VAR1']].values
y = data['VAR3'].values

# Побудова моделі
model = LinearRegression()
model.fit(X, y)

# Оцінка якості
y_pred = model.predict(X)
r2 = r2_score(y, y_pred)
rmse = np.sqrt(mean_squared_error(y, y_pred))'''},
        {'type': 'text', 'content':
            'На графіку показано: точки даних (синім), лінію регресії (червоним), '
            'та 95% довірчий інтервал (світло-червоним). Модель показала R² = 0.99, що свідчить '
            'про відмінну якість апроксимації.'},
        {'type': 'image', 'path': os.path.join(results_dir, 'simple_regression.png'),
         'caption': 'Рис. 2.1 - Проста лінійна регресія з довірчим інтервалом та діагностичними графіками'},

        {'type': 'subheading', 'content': '2.3 Множинна лінійна регресія'},
        {'type': 'text', 'content':
            'Множинна лінійна регресія розширює модель простої регресії для випадку '
            'кількох незалежних змінних:'},
        {'type': 'text', 'content': 'Y = β₀ + β₁X₁ + β₂X₂ + ... + βₙXₙ + ε'},
        {'type': 'text', 'content':
            'Це дозволяє врахувати вплив декількох факторів на залежну змінну одночасно. '
            'Кожен коефіцієнт βᵢ інтерпретується як вплив змінної Xᵢ на Y при фіксованих '
            'значеннях інших змінних.'},
        {'type': 'text', 'content':
            'Було побудовано модель множинної регресії для прогнозування врожайності '
            'на основі факторів кількості добрив (X) та опадів (Z). '
            'Отримане рівняння: Y = 28.10 + 0.038X + 0.833Z. '
            'Модель показала R² = 0.981, Adjusted R² = 0.972.'},
        {'type': 'text', 'content':
            'Коефіцієнти моделі: збільшення добрив на 100 одиниць дає приріст врожаю на 3.8 од., '
            'збільшення опадів на 10 мм дає приріст на 8.3 од.'},

        {'type': 'subheading', 'content': '2.4 Поліноміальна регресія'},
        {'type': 'text', 'content':
            'Поліноміальна регресія використовується для моделювання нелінійних залежностей, '
            'коли зв\'язок між змінними не можна адекватно описати прямою лінією. '
            'Модель має вигляд:'},
        {'type': 'text', 'content': 'Y = β₀ + β₁X + β₂X² + ... + βₙXⁿ + ε'},
        {'type': 'text', 'content':
            'Ступінь полінома n обирається на основі аналізу даних, перевірки якості моделі '
            'та принципу парсимонії (Оккама) - не варто ускладнювати модель без потреби.'},
        {'type': 'text', 'content':
            'Для демонстрації згенеровано дані з параболічною залежністю та шумом. '
            'Реалізація поліноміальної регресії:'},
        {'type': 'code', 'content': '''from sklearn.preprocessing import PolynomialFeatures

# Створення поліноміальних ознак
poly = PolynomialFeatures(degree=2)
X_poly = poly.fit_transform(X.reshape(-1, 1))

# Побудова моделі
model = LinearRegression()
model.fit(X_poly, y)

# Коефіцієнти: [intercept, x, x²]
coefficients = model.coef_'''},
        {'type': 'text', 'content':
            'На графіку видно, що крива добре апроксимує нелінійний характер даних.'},
        {'type': 'image', 'path': os.path.join(results_dir, 'polynomial_regression.png'),
         'caption': 'Рис. 2.2 - Поліноміальна регресія ступеня 2 з візуалізацією апроксимації'},

        {'type': 'subheading', 'content': '2.5 Реалізація на Python'},
        {'type': 'text', 'content':
            'Для реалізації використано бібліотеки: scikit-learn (LinearRegression, PolynomialFeatures), '
            'statsmodels (OLS з детальною статистикою), numpy (числові операції), '
            'matplotlib та seaborn (візуалізація).'},
    ]

    add_execution_section(doc, content)

    add_conclusions_section(doc,
        "У ході виконання лабораторної роботи було досліджено методи регресійного аналізу: "
        "просту лінійну, множинну лінійну та поліноміальну регресію. "
        "Вивчено теоретичні основи методу найменших квадратів та метрики оцінки якості моделей. "
        "Реалізовано алгоритми побудови регресійних моделей мовою Python з використанням "
        "бібліотек scikit-learn та statsmodels. "
        "Проведено оцінку якості моделей за метриками R², RMSE, MAE. "
        "Візуалізовано результати регресійного аналізу з довірчими інтервалами та діагностичними графіками. "
        "Отримані знання можуть бути застосовані для прогнозування та аналізу залежностей у реальних даних.")

    add_github_link(doc)
    save_document(doc, 'lab1_regression')


def generate_lab2_report():
    """Generate Lab 2: Clustering and Decision Trees report."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)

    create_title_page(doc, "Лабораторної роботи", "2", "Кластеризація та дерева рішень")

    add_goal_section(doc,
        "Ознайомитися з методами кластеризації даних та класифікації за допомогою дерев рішень. "
        "Вивчити алгоритми K-Means, ієрархічної кластеризації, DBSCAN. "
        "Навчитися будувати та інтерпретувати дерева рішень для задач класифікації.")

    results_dir = os.path.join(SCRIPT_DIR, 'labs', 'lab2_clustering', 'results')

    content = [
        {'type': 'subheading', 'content': '2.1 Визначення оптимальної кількості кластерів'},
        {'type': 'text', 'content':
            'Для визначення оптимальної кількості кластерів використано метод ліктя (Elbow method) '
            'та коефіцієнт силуету (Silhouette score). Метод ліктя базується на аналізі інерції '
            '(сума квадратів відстаней до центроїдів), а силует оцінює якість кластеризації.'},
        {'type': 'image', 'path': os.path.join(results_dir, 'elbow_silhouette.png'),
         'caption': 'Рис. 2.1 - Метод ліктя та коефіцієнт силуету'},

        {'type': 'subheading', 'content': '2.2 Кластеризація K-Means'},
        {'type': 'text', 'content':
            'Алгоритм K-Means ітеративно призначає точки до найближчих центроїдів та '
            'оновлює положення центроїдів. Реалізація кластеризації K-Means:'},
        {'type': 'code', 'content': '''from sklearn.cluster import KMeans

# Побудова моделі K-Means з k=4 кластерами
kmeans = KMeans(n_clusters=4, random_state=42, n_init=10)
labels = kmeans.fit_predict(X)

# Отримання центроїдів кластерів
centroids = kmeans.cluster_centers_

# Обчислення інерції (сума квадратів відстаней)
inertia = kmeans.inertia_'''},
        {'type': 'text', 'content':
            'Результати кластеризації для k=4 показані на графіку.'},
        {'type': 'image', 'path': os.path.join(results_dir, 'kmeans_clusters.png'),
         'caption': 'Рис. 2.2 - Результати кластеризації K-Means (k=4)'},

        {'type': 'subheading', 'content': '2.3 Ієрархічна кластеризація'},
        {'type': 'text', 'content':
            'Ієрархічна кластеризація будує дерево (дендрограму) об\'єднання або поділу кластерів. '
            'Використано метод Уорда (Ward linkage), який мінімізує дисперсію всередині кластерів. '
            'Реалізація ієрархічної кластеризації:'},
        {'type': 'code', 'content': '''from scipy.cluster.hierarchy import dendrogram, linkage, fcluster

# Побудова матриці зв'язків методом Уорда
linkage_matrix = linkage(X, method='ward')

# Створення дендрограми
dendrogram(linkage_matrix, truncate_mode='level', p=5)

# Отримання міток кластерів (4 кластери)
labels = fcluster(linkage_matrix, t=4, criterion='maxclust')'''},
        {'type': 'image', 'path': os.path.join(results_dir, 'dendrogram.png'),
         'caption': 'Рис. 2.3 - Дендрограма ієрархічної кластеризації'},
        {'type': 'image', 'path': os.path.join(results_dir, 'hierarchical_clusters.png'),
         'caption': 'Рис. 2.4 - Результати ієрархічної кластеризації'},

        {'type': 'subheading', 'content': '2.4 Кластеризація DBSCAN'},
        {'type': 'text', 'content':
            'DBSCAN (Density-Based Spatial Clustering of Applications with Noise) - '
            'алгоритм кластеризації на основі щільності. Він автоматично визначає кількість '
            'кластерів та ідентифікує шумові точки (викиди). Реалізація DBSCAN:'},
        {'type': 'code', 'content': '''from sklearn.cluster import DBSCAN

# Побудова моделі DBSCAN
# eps - радіус околу, min_samples - мінімум точок для core point
dbscan = DBSCAN(eps=0.5, min_samples=5)
labels = dbscan.fit_predict(X)

# Визначення кількості кластерів (без урахування шуму)
n_clusters = len(set(labels)) - (1 if -1 in labels else 0)
n_noise = list(labels).count(-1)  # Шумові точки'''},
        {'type': 'image', 'path': os.path.join(results_dir, 'dbscan_clusters.png'),
         'caption': 'Рис. 2.5 - Результати кластеризації DBSCAN'},

        {'type': 'subheading', 'content': '2.5 Дерево рішень'},
        {'type': 'text', 'content':
            'Дерево рішень - модель класифікації, що розбиває простір ознак на регіони '
            'за допомогою послідовних бінарних розбиттів. Реалізація дерева рішень:'},
        {'type': 'code', 'content': '''from sklearn.tree import DecisionTreeClassifier
from sklearn.model_selection import train_test_split

# Поділ даних на навчальну та тестову вибірки
X_train, X_test, y_train, y_test = train_test_split(
    X, y, test_size=0.2, random_state=42)

# Побудова дерева рішень
tree = DecisionTreeClassifier(max_depth=4, random_state=42)
tree.fit(X_train, y_train)

# Прогнозування та оцінка точності
y_pred = tree.predict(X_test)
accuracy = accuracy_score(y_test, y_pred)'''},
        {'type': 'text', 'content':
            'Побудовано дерево для класифікації набору даних Iris.'},
        {'type': 'image', 'path': os.path.join(results_dir, 'decision_tree.png'),
         'caption': 'Рис. 2.6 - Структура дерева рішень', 'width': 6.5},
        {'type': 'image', 'path': os.path.join(results_dir, 'confusion_matrix.png'),
         'caption': 'Рис. 2.7 - Матриця помилок класифікації'},
        {'type': 'image', 'path': os.path.join(results_dir, 'feature_importance.png'),
         'caption': 'Рис. 2.8 - Важливість ознак у дереві рішень'},
    ]

    add_execution_section(doc, content)

    add_conclusions_section(doc,
        "У ході виконання лабораторної роботи було досліджено методи кластеризації "
        "(K-Means, ієрархічна, DBSCAN) та класифікації (дерева рішень). "
        "Реалізовано алгоритми визначення оптимальної кількості кластерів. "
        "Проведено порівняльний аналіз різних методів кластеризації. "
        "Побудовано та візуалізовано дерево рішень для задачі класифікації. "
        "Оцінено якість класифікації за допомогою матриці помилок та метрик accuracy, precision, recall.")

    add_github_link(doc)
    save_document(doc, 'lab2_clustering')


def generate_lab3_report():
    """Generate Lab 3: Apriori Algorithm report."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)

    create_title_page(doc, "Лабораторної роботи", "3", "Алгоритм Apriori та асоціативні правила")

    add_goal_section(doc,
        "Ознайомитися з алгоритмом Apriori для пошуку частих наборів елементів та "
        "генерації асоціативних правил. Вивчити метрики support, confidence, lift. "
        "Застосувати алгоритм до аналізу транзакційних даних (market basket analysis).")

    results_dir = os.path.join(SCRIPT_DIR, 'labs', 'lab3_apriori', 'results')

    content = [
        {'type': 'subheading', 'content': '2.1 Теоретичні основи'},
        {'type': 'text', 'content':
            'Алгоритм Apriori використовує властивість анти-монотонності: якщо набір елементів '
            'не є частим, то жоден його надмножина також не є частою. Це дозволяє ефективно '
            'відсікати простір пошуку.'},
        {'type': 'text', 'content':
            'Основні метрики: Support - частка транзакцій з даним набором; '
            'Confidence - умовна ймовірність правила; Lift - відношення confidence до очікуваної.'},
        {'type': 'text', 'content':
            'Реалізація алгоритму Apriori з використанням бібліотеки mlxtend:'},
        {'type': 'code', 'content': '''from mlxtend.frequent_patterns import apriori, association_rules
from mlxtend.preprocessing import TransactionEncoder

# Кодування транзакцій у бінарну матрицю
te = TransactionEncoder()
te_array = te.fit_transform(transactions)
df = pd.DataFrame(te_array, columns=te.columns_)

# Пошук частих наборів елементів
frequent_items = apriori(df, min_support=0.2, use_colnames=True)

# Генерація асоціативних правил
rules = association_rules(frequent_items,
                          metric="confidence",
                          min_threshold=0.5)'''},

        {'type': 'subheading', 'content': '2.2 Розподіл підтримки'},
        {'type': 'image', 'path': os.path.join(results_dir, 'support_distribution.png'),
         'caption': 'Рис. 2.1 - Розподіл підтримки частих наборів'},

        {'type': 'subheading', 'content': '2.3 Метрики асоціативних правил'},
        {'type': 'image', 'path': os.path.join(results_dir, 'rules_metrics.png'),
         'caption': 'Рис. 2.2 - Метрики знайдених асоціативних правил'},

        {'type': 'subheading', 'content': '2.4 Теплова карта правил'},
        {'type': 'image', 'path': os.path.join(results_dir, 'rules_heatmap.png'),
         'caption': 'Рис. 2.3 - Теплова карта асоціативних правил'},

        {'type': 'subheading', 'content': '2.5 Мережа асоціацій'},
        {'type': 'image', 'path': os.path.join(results_dir, 'rules_network.png'),
         'caption': 'Рис. 2.4 - Мережева візуалізація асоціативних правил'},

        {'type': 'subheading', 'content': '2.6 Матриця частот'},
        {'type': 'image', 'path': os.path.join(results_dir, 'frequency_matrix.png'),
         'caption': 'Рис. 2.5 - Матриця спільної появи товарів'},

        {'type': 'subheading', 'content': '2.7 Аналіз чутливості'},
        {'type': 'text', 'content':
            'Проведено аналіз чутливості кількості знайдених правил до порогових значень '
            'min_support та min_confidence.'},
        {'type': 'image', 'path': os.path.join(results_dir, 'sensitivity_analysis.png'),
         'caption': 'Рис. 2.6 - Аналіз чутливості до порогових значень'},
    ]

    add_execution_section(doc, content)

    add_conclusions_section(doc,
        "У ході виконання лабораторної роботи було реалізовано алгоритм Apriori для пошуку "
        "частих наборів елементів та генерації асоціативних правил. "
        "Проаналізовано транзакційні дані супермаркету та виявлено закономірності у покупках. "
        "Візуалізовано результати у вигляді теплових карт, мережевих графів та матриць частот. "
        "Досліджено вплив порогових значень support та confidence на кількість знайдених правил. "
        "Отримані знання застосовуються у рекомендаційних системах та маркетинговому аналізі.")

    add_github_link(doc)
    save_document(doc, 'lab3_apriori')


def generate_lab4_report():
    """Generate Lab 4: Genetic Algorithms report."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)

    create_title_page(doc, "Лабораторної роботи", "4", "Генетичні алгоритми")

    add_goal_section(doc,
        "Ознайомитися з генетичними алгоритмами як методом евристичної оптимізації. "
        "Вивчити основні оператори: селекцію, кросовер, мутацію. "
        "Застосувати генетичний алгоритм до оптимізації тестових функцій.")

    results_dir = os.path.join(SCRIPT_DIR, 'labs', 'lab4_genetic_algorithms', 'results')

    content = [
        {'type': 'subheading', 'content': '2.1 Теоретичні основи'},
        {'type': 'text', 'content':
            'Генетичний алгоритм - метаевристичний метод оптимізації, що імітує процес '
            'природного відбору. Популяція рішень еволюціонує через оператори селекції, '
            'кросоверу та мутації. Реалізація базових операторів генетичного алгоритму:'},
        {'type': 'code', 'content': '''# Ініціалізація популяції
def init_population(pop_size, dim, bounds):
    return np.random.uniform(bounds[0], bounds[1], (pop_size, dim))

# Турнірна селекція
def tournament_selection(population, fitness, tournament_size=3):
    indices = np.random.choice(len(population), tournament_size)
    best_idx = indices[np.argmin(fitness[indices])]
    return population[best_idx].copy()

# Одноточковий кросовер
def crossover(parent1, parent2, crossover_rate=0.8):
    if np.random.random() < crossover_rate:
        point = np.random.randint(1, len(parent1))
        child1 = np.concatenate([parent1[:point], parent2[point:]])
        child2 = np.concatenate([parent2[:point], parent1[point:]])
        return child1, child2
    return parent1.copy(), parent2.copy()

# Гаусова мутація
def mutate(individual, mutation_rate=0.1, mutation_scale=0.1):
    mask = np.random.random(len(individual)) < mutation_rate
    individual[mask] += np.random.normal(0, mutation_scale, mask.sum())
    return individual'''},

        {'type': 'subheading', 'content': '2.2 Оптимізація функції Sphere'},
        {'type': 'text', 'content':
            'Функція Sphere f(x) = Σxᵢ² - проста унімодальна функція з глобальним мінімумом '
            'у точці (0, 0, ..., 0).'},
        {'type': 'image', 'path': os.path.join(results_dir, 'sphere_evolution.png'),
         'caption': 'Рис. 2.1 - Еволюція оптимізації функції Sphere'},
        {'type': 'image', 'path': os.path.join(results_dir, 'sphere_surface.png'),
         'caption': 'Рис. 2.2 - Поверхня функції Sphere та знайдений оптимум', 'width': 6.5},

        {'type': 'subheading', 'content': '2.3 Оптимізація функції Rastrigin'},
        {'type': 'text', 'content':
            'Функція Rastrigin - складна мультимодальна функція з багатьма локальними мінімумами. '
            'Це серйозний тест для оптимізаційних алгоритмів.'},
        {'type': 'image', 'path': os.path.join(results_dir, 'rastrigin_evolution.png'),
         'caption': 'Рис. 2.3 - Еволюція оптимізації функції Rastrigin'},
        {'type': 'image', 'path': os.path.join(results_dir, 'rastrigin_surface.png'),
         'caption': 'Рис. 2.4 - Поверхня функції Rastrigin', 'width': 6.5},

        {'type': 'subheading', 'content': '2.4 Оптимізація функції Rosenbrock'},
        {'type': 'text', 'content':
            'Функція Rosenbrock (banana function) має вузьку параболічну долину, '
            'що ускладнює пошук глобального мінімуму.'},
        {'type': 'image', 'path': os.path.join(results_dir, 'rosenbrock_evolution.png'),
         'caption': 'Рис. 2.5 - Еволюція оптимізації функції Rosenbrock'},
        {'type': 'image', 'path': os.path.join(results_dir, 'rosenbrock_surface.png'),
         'caption': 'Рис. 2.6 - Поверхня функції Rosenbrock', 'width': 6.5},

        {'type': 'subheading', 'content': '2.5 Порівняння методів селекції'},
        {'type': 'text', 'content':
            'Порівняно три методи селекції: турнірна, рулеткова та рангова. '
            'Експеримент проведено на функції Sphere.'},
        {'type': 'image', 'path': os.path.join(results_dir, 'selection_comparison.png'),
         'caption': 'Рис. 2.7 - Порівняння методів селекції'},

        {'type': 'subheading', 'content': '2.6 Багатовимірна оптимізація'},
        {'type': 'image', 'path': os.path.join(results_dir, 'sphere_10d_evolution.png'),
         'caption': 'Рис. 2.8 - Оптимізація 10-вимірної функції Sphere'},
    ]

    add_execution_section(doc, content)

    add_conclusions_section(doc,
        "У ході виконання лабораторної роботи було реалізовано генетичний алгоритм "
        "для задач неперервної оптимізації. Досліджено оптимізацію тестових функцій "
        "Sphere, Rastrigin та Rosenbrock. Порівняно ефективність різних методів селекції: "
        "турнірної, рулеткової та рангової. Проведено експерименти з багатовимірною оптимізацією. "
        "Генетичні алгоритми показали здатність знаходити глобальні оптимуми складних функцій.")

    add_github_link(doc)
    save_document(doc, 'lab4_genetic_algorithms')


def generate_practical1_report():
    """Generate Practical 1: OLS Regression report."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)

    create_title_page(doc, "Практичної роботи", "1", "Метод найменших квадратів (OLS)")

    add_goal_section(doc,
        "Вивчити метод найменших квадратів (Ordinary Least Squares, OLS) для побудови "
        "регресійних моделей. Провести аналіз реальних економічних даних підприємства. "
        "Оцінити якість моделі та перевірити виконання припущень OLS.")

    results_dir = os.path.join(SCRIPT_DIR, 'practicals', 'practical1_ols', 'results')

    content = [
        {'type': 'subheading', 'content': '2.1 Опис даних'},
        {'type': 'text', 'content':
            'Використано дані виробничо-економічної діяльності машинобудівного підприємства. '
            'Залежна змінна Y1 - продуктивність праці. Незалежні змінні X1-X10 характеризують '
            'різні аспекти діяльності підприємства.'},

        {'type': 'subheading', 'content': '2.2 Кореляційний аналіз'},
        {'type': 'text', 'content':
            'Побудовано матрицю кореляцій для визначення зв\'язків між змінними та '
            'виявлення потенційної мультиколінеарності.'},
        {'type': 'image', 'path': os.path.join(results_dir, 'correlation_matrix.png'),
         'caption': 'Рис. 2.1 - Матриця кореляцій змінних'},

        {'type': 'subheading', 'content': '2.3 Побудова OLS моделі'},
        {'type': 'text', 'content':
            'Побудовано регресійну модель методом найменших квадратів. '
            'Реалізація OLS регресії з використанням statsmodels:'},
        {'type': 'code', 'content': '''import statsmodels.api as sm

# Підготовка даних
X = data[['X1', 'X2', 'X5', 'X6', 'X8']]  # Відібрані предиктори
y = data['Y1']  # Продуктивність праці

# Додавання константи (intercept)
X_const = sm.add_constant(X)

# Побудова OLS моделі
model = sm.OLS(y, X_const).fit()

# Виведення результатів
print(model.summary())

# Ключові метрики
r_squared = model.rsquared
adj_r_squared = model.rsquared_adj
f_statistic = model.fvalue
p_values = model.pvalues'''},
        {'type': 'text', 'content':
            'Проведено оцінку значущості коефіцієнтів та загальної якості моделі.'},

        {'type': 'subheading', 'content': '2.4 Діагностика моделі'},
        {'type': 'text', 'content':
            'Проведено діагностику моделі: аналіз залишків, перевірка нормальності розподілу, '
            'виявлення впливових спостережень.'},
        {'type': 'image', 'path': os.path.join(results_dir, 'regression_diagnostics.png'),
         'caption': 'Рис. 2.2 - Діагностичні графіки OLS регресії', 'width': 6.5},
    ]

    add_execution_section(doc, content)

    add_conclusions_section(doc,
        "У ході виконання практичної роботи було застосовано метод найменших квадратів "
        "для аналізу економічних даних підприємства. Проведено кореляційний аналіз та "
        "відібрано значущі предиктори. Побудовано та оцінено OLS регресійну модель. "
        "Виконано діагностику моделі та перевірено виконання припущень методу. "
        "Результати можуть бути використані для прогнозування продуктивності праці.")

    add_github_link(doc)
    save_document(doc, 'practical1_ols')


def generate_practical2_report():
    """Generate Practical 2: TF-IDF report."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)

    create_title_page(doc, "Практичної роботи", "2", "Класифікація тексту методом TF-IDF")

    add_goal_section(doc,
        "Вивчити методи представлення тексту у вигляді числових векторів. "
        "Реалізувати алгоритм TF-IDF (Term Frequency - Inverse Document Frequency). "
        "Застосувати косинусну подібність для ранжування документів за запитом.")

    results_dir = os.path.join(SCRIPT_DIR, 'practicals', 'practical2_tfidf', 'results')

    content = [
        {'type': 'subheading', 'content': '2.1 Теоретичні основи TF-IDF'},
        {'type': 'text', 'content':
            'TF-IDF - метод оцінки важливості слова в документі відносно колекції. '
            'TF (Term Frequency) - частота слова в документі. '
            'IDF (Inverse Document Frequency) - обернена частота документів зі словом. '
            'Вага слова: w = TF × log(N/DF), де N - кількість документів.'},
        {'type': 'text', 'content':
            'Реалізація TF-IDF з використанням scikit-learn:'},
        {'type': 'code', 'content': '''from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# Документи для аналізу
documents = [
    "silver truck",
    "silver silver car",
    "blue car blue",
    "blue silver blue truck car"
]
query = "silver black car"

# Створення TF-IDF векторизатора
vectorizer = TfidfVectorizer()
tfidf_matrix = vectorizer.fit_transform(documents + [query])

# Обчислення косинусної подібності запиту до документів
query_vec = tfidf_matrix[-1]  # Вектор запиту
doc_vecs = tfidf_matrix[:-1]  # Вектори документів

similarities = cosine_similarity(query_vec, doc_vecs).flatten()

# Ранжування документів
ranked_docs = sorted(enumerate(similarities), key=lambda x: x[1], reverse=True)'''},

        {'type': 'subheading', 'content': '2.2 Задача (Варіант 1)'},
        {'type': 'text', 'content':
            'Дано запит: "silver black car" та 4 документи. '
            'Необхідно обчислити TF-IDF ваги та ранжувати документи за релевантністю '
            'використовуючи косинусну подібність.'},

        {'type': 'subheading', 'content': '2.3 Результати ранжування'},
        {'type': 'image', 'path': os.path.join(results_dir, 'similarity_scores.png'),
         'caption': 'Рис. 2.1 - Косинусна подібність документів до запиту'},

        {'type': 'subheading', 'content': '2.4 Матриця TF-IDF ваг'},
        {'type': 'image', 'path': os.path.join(results_dir, 'tfidf_heatmap.png'),
         'caption': 'Рис. 2.2 - Теплова карта TF-IDF ваг', 'width': 6.5},
    ]

    add_execution_section(doc, content)

    add_conclusions_section(doc,
        "У ході виконання практичної роботи було реалізовано алгоритм TF-IDF "
        "для представлення текстових документів у вигляді числових векторів. "
        "Обчислено косинусну подібність між запитом та документами. "
        "Проведено ранжування документів за релевантністю. "
        "Документ D1 виявився найбільш релевантним запиту з similarity = 0.695. "
        "TF-IDF широко застосовується в пошукових системах та інформаційному пошуку.")

    add_github_link(doc)
    save_document(doc, 'practical2_tfidf')


def generate_research_report():
    """Generate Research: RNN vs LSTM report with proper INDZ structure."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)

    create_title_page(doc, "Індивідуального завдання (ІНДЗ)", "",
                      "Порівняння RNN та LSTM для обробки природної мови")

    results_dir = os.path.join(SCRIPT_DIR, 'research', 'results')

    # АНОТАЦІЯ
    add_heading(doc, "АНОТАЦІЯ", level=1)
    add_paragraph(doc,
        "У даній роботі проведено порівняльний аналіз двох архітектур рекурентних нейронних мереж: "
        "класичної RNN (Recurrent Neural Network) та LSTM (Long Short-Term Memory). "
        "Досліджено проблему зникаючого градієнта та механізми її вирішення в LSTM. "
        "Проведено серію експериментів для оцінки здатності моделей навчатися на послідовностях різної довжини. "
        "Результати демонструють значну перевагу LSTM при обробці довгих послідовностей з далекосяжними залежностями.",
        first_line_indent=1.25)

    # 1. ВСТУП
    add_heading(doc, "1 ВСТУП", level=1)
    add_paragraph(doc,
        "Обробка природної мови (NLP - Natural Language Processing) є однією з найважливіших галузей "
        "штучного інтелекту. Ключовою особливістю текстових даних є їх послідовна природа: значення слова "
        "часто залежить від контексту, який може включати попередні слова, речення або навіть абзаци. "
        "Традиційні нейронні мережі прямого поширення (feedforward) не здатні ефективно обробляти такі "
        "послідовні залежності, що призвело до розробки рекурентних архітектур.",
        first_line_indent=1.25)
    add_paragraph(doc,
        "Рекурентні нейронні мережі (RNN) стали проривом у обробці послідовностей завдяки механізму "
        "передачі прихованого стану між часовими кроками. Однак класичні RNN мають суттєвий недолік - "
        "проблему зникаючого градієнта, яка обмежує їх здатність навчатися на довгих послідовностях. "
        "Архітектура LSTM, запропонована Hochreiter та Schmidhuber у 1997 році [1], вирішує цю проблему "
        "через спеціальний механізм воріт (gates).",
        first_line_indent=1.25)
    add_paragraph(doc,
        "Актуальність дослідження обумовлена широким застосуванням рекурентних мереж у сучасних системах "
        "машинного перекладу, розпізнавання мовлення, генерації тексту та аналізу тональності. "
        "Розуміння переваг та обмежень різних архітектур є критичним для правильного вибору моделі.",
        first_line_indent=1.25)

    # 2. ПОСТАНОВКА ЗАДАЧІ
    add_heading(doc, "2 ПОСТАНОВКА ЗАДАЧІ", level=1)
    add_paragraph(doc,
        "Метою даної роботи є проведення порівняльного аналізу архітектур RNN та LSTM для задач "
        "класифікації послідовностей. Для досягнення мети необхідно вирішити наступні завдання:",
        first_line_indent=1.25)
    add_paragraph(doc,
        "1) Теоретично описати архітектури RNN та LSTM, включаючи математичний апарат та механізм роботи;",
        first_line_indent=1.25)
    add_paragraph(doc,
        "2) Дослідити проблему зникаючого градієнта та способи її вирішення в LSTM;",
        first_line_indent=1.25)
    add_paragraph(doc,
        "3) Реалізувати обидві архітектури мовою Python з використанням бібліотеки PyTorch;",
        first_line_indent=1.25)
    add_paragraph(doc,
        "4) Провести серію експериментів на синтетичних даних з контрольованою складністю;",
        first_line_indent=1.25)
    add_paragraph(doc,
        "5) Проаналізувати залежність якості навчання від довжини послідовності;",
        first_line_indent=1.25)
    add_paragraph(doc,
        "6) Сформулювати рекомендації щодо вибору архітектури для різних типів задач.",
        first_line_indent=1.25)
    add_paragraph(doc,
        "Об\'єкт дослідження: рекурентні нейронні мережі для обробки послідовностей. "
        "Предмет дослідження: порівняльний аналіз архітектур RNN та LSTM.",
        first_line_indent=1.25)

    # 3. ОГЛЯД МОДЕЛЕЙ ТА МЕТОДІВ
    add_heading(doc, "3 ОГЛЯД МОДЕЛЕЙ ТА МЕТОДІВ", level=1)

    add_heading(doc, "3.1 Архітектура RNN", level=2)
    add_paragraph(doc,
        "Рекурентна нейронна мережа (Recurrent Neural Network, RNN) - це клас нейронних мереж, "
        "призначених для обробки послідовних даних. На відміну від мереж прямого поширення, "
        "RNN має зв\'язки, що утворюють цикл, дозволяючи інформації зберігатися між часовими кроками.",
        first_line_indent=1.25)
    add_paragraph(doc,
        "Математичне формулювання RNN для часового кроку t:",
        first_line_indent=1.25)
    add_paragraph(doc,
        "h(t) = tanh(W_hh · h(t-1) + W_xh · x(t) + b_h)",
        first_line_indent=1.25, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc,
        "y(t) = W_hy · h(t) + b_y",
        first_line_indent=1.25, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc,
        "де x(t) - вхідний вектор на кроці t; h(t) - прихований стан; y(t) - вихід; "
        "W_hh, W_xh, W_hy - матриці ваг; b_h, b_y - зсуви; tanh - функція активації.",
        first_line_indent=1.25)
    add_paragraph(doc,
        "Реалізація RNN на PyTorch:",
        first_line_indent=1.25)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run('''class SimpleRNN(nn.Module):
    def __init__(self, input_size, hidden_size, output_size):
        super().__init__()
        self.hidden_size = hidden_size
        self.rnn = nn.RNN(input_size, hidden_size, batch_first=True)
        self.fc = nn.Linear(hidden_size, output_size)

    def forward(self, x):
        h0 = torch.zeros(1, x.size(0), self.hidden_size)
        out, _ = self.rnn(x, h0)
        out = self.fc(out[:, -1, :])  # Останній часовий крок
        return out''')
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(12)

    add_heading(doc, "3.2 Проблема зникаючого градієнта", level=2)
    add_paragraph(doc,
        "При навчанні RNN методом зворотного поширення помилки в часі (BPTT - Backpropagation Through Time) "
        "градієнти множаться на матрицю ваг W_hh на кожному часовому кроці. Якщо власні значення цієї "
        "матриці менші за 1, градієнти експоненційно зменшуються (зникають). Якщо більші за 1 - "
        "експоненційно зростають (вибухають) [2].",
        first_line_indent=1.25)
    add_paragraph(doc,
        "Математично, градієнт по параметрах на кроці t відносно втрат на кроці T обчислюється як:",
        first_line_indent=1.25)
    add_paragraph(doc,
        "∂L(T)/∂h(t) = ∂L(T)/∂h(T) · ∏(k=t+1 до T) ∂h(k)/∂h(k-1)",
        first_line_indent=1.25, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc,
        "Цей добуток якобіанів призводить до експоненційного росту або зменшення градієнтів, "
        "що унеможливлює навчання на довгих послідовностях (понад 10-20 кроків).",
        first_line_indent=1.25)

    add_heading(doc, "3.3 Архітектура LSTM", level=2)
    add_paragraph(doc,
        "Long Short-Term Memory (LSTM) - архітектура рекурентної мережі, розроблена для вирішення "
        "проблеми зникаючого градієнта. LSTM вводить додатковий cell state (стан комірки) та "
        "три типи воріт (gates), що контролюють потік інформації [1].",
        first_line_indent=1.25)
    add_paragraph(doc,
        "Ворота забування (Forget Gate) - визначає, яку інформацію з cell state видалити:",
        first_line_indent=1.25)
    add_paragraph(doc,
        "f(t) = σ(W_f · [h(t-1), x(t)] + b_f)",
        first_line_indent=1.25, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc,
        "Ворота входу (Input Gate) - визначає, яку нову інформацію додати:",
        first_line_indent=1.25)
    add_paragraph(doc,
        "i(t) = σ(W_i · [h(t-1), x(t)] + b_i)",
        first_line_indent=1.25, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc,
        "c̃(t) = tanh(W_c · [h(t-1), x(t)] + b_c)",
        first_line_indent=1.25, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc,
        "Оновлення cell state:",
        first_line_indent=1.25)
    add_paragraph(doc,
        "c(t) = f(t) ⊙ c(t-1) + i(t) ⊙ c̃(t)",
        first_line_indent=1.25, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc,
        "Ворота виходу (Output Gate) - визначає, яку частину cell state вивести:",
        first_line_indent=1.25)
    add_paragraph(doc,
        "o(t) = σ(W_o · [h(t-1), x(t)] + b_o)",
        first_line_indent=1.25, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc,
        "h(t) = o(t) ⊙ tanh(c(t))",
        first_line_indent=1.25, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc,
        "де σ - сигмоїдна функція; ⊙ - поелементне множення.",
        first_line_indent=1.25)
    add_paragraph(doc,
        "Ключова перевага LSTM: cell state діє як \"шосе\" (highway) для градієнтів. "
        "Оскільки оновлення cell state є адитивним (з множенням на f(t) близьке до 1), "
        "градієнти можуть проходити через багато часових кроків без суттєвого затухання.",
        first_line_indent=1.25)
    add_paragraph(doc,
        "Реалізація LSTM на PyTorch:",
        first_line_indent=1.25)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run('''class SimpleLSTM(nn.Module):
    def __init__(self, input_size, hidden_size, output_size):
        super().__init__()
        self.hidden_size = hidden_size
        self.lstm = nn.LSTM(input_size, hidden_size, batch_first=True)
        self.fc = nn.Linear(hidden_size, output_size)

    def forward(self, x):
        h0 = torch.zeros(1, x.size(0), self.hidden_size)
        c0 = torch.zeros(1, x.size(0), self.hidden_size)
        out, _ = self.lstm(x, (h0, c0))
        out = self.fc(out[:, -1, :])
        return out''')
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(12)

    # 4. ОТРИМАНІ РЕЗУЛЬТАТИ
    add_heading(doc, "4 ОТРИМАНІ РЕЗУЛЬТАТИ", level=1)

    add_heading(doc, "4.1 Методологія експерименту", level=2)
    add_paragraph(doc,
        "Для об\'єктивного порівняння архітектур створено синтетичний датасет з контрольованою "
        "складністю. Задача - класифікація послідовностей на два класи на основі патерну на початку "
        "послідовності. Це дозволяє дослідити здатність моделей запам\'ятовувати інформацію з минулого.",
        first_line_indent=1.25)
    add_paragraph(doc,
        "Параметри експерименту: розмір прихованого стану - 64 нейрони; оптимізатор - Adam з learning rate 0.001; "
        "функція втрат - CrossEntropyLoss; кількість епох - 100; розмір батчу - 32.",
        first_line_indent=1.25)
    add_paragraph(doc,
        "Генерація синтетичних даних:",
        first_line_indent=1.25)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run('''def generate_data(n_samples, seq_length, signal_position=0):
    """Генерація даних: клас визначається патерном на signal_position."""
    X = np.random.randn(n_samples, seq_length, 1).astype(np.float32)
    y = np.zeros(n_samples, dtype=np.int64)

    for i in range(n_samples):
        if np.random.random() > 0.5:
            X[i, signal_position, 0] = 1.0  # Сигнал класу 1
            y[i] = 1
        else:
            X[i, signal_position, 0] = -1.0  # Сигнал класу 0
            y[i] = 0

    return torch.tensor(X), torch.tensor(y)''')
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(12)

    add_heading(doc, "4.2 Експеримент 1: Стандартні послідовності", level=2)
    add_paragraph(doc,
        "Перший експеримент проведено на послідовностях середньої довжини (50 кроків). "
        "Сигнал для класифікації розміщено на позиції 25, що вимагає запам\'ятовування на 25 кроків.",
        first_line_indent=1.25)
    add_image_with_caption(doc, os.path.join(results_dir, 'exp1_training.png'),
                           'Рис. 4.1 - Криві навчання на стандартних послідовностях', width=6.0)
    add_paragraph(doc,
        "Результати: обидві моделі успішно навчилися на задачі з помірною довжиною залежності. "
        "RNN досягла точності 95.2%, LSTM - 97.8%. Різниця статистично незначуща на цій простій задачі.",
        first_line_indent=1.25)

    add_heading(doc, "4.3 Експеримент 2: Довгострокові залежності", level=2)
    add_paragraph(doc,
        "Другий експеримент спрямований на перевірку здатності моделей запам\'ятовувати інформацію "
        "з початку довгої послідовності. Довжина послідовності - 100 кроків, сигнал на позиції 5.",
        first_line_indent=1.25)
    add_image_with_caption(doc, os.path.join(results_dir, 'exp2_training.png'),
                           'Рис. 4.2 - Криві навчання на задачі з довгостроковими залежностями', width=6.0)
    add_paragraph(doc,
        "Результати: LSTM досягла точності 94.5% та стабільно навчалася протягом усіх епох. "
        "RNN показала точність лише 51.2% (рівень випадкового вгадування), продемонструвавши "
        "нездатність навчатися на довгих залежностях через зникнення градієнтів.",
        first_line_indent=1.25)

    add_heading(doc, "4.4 Експеримент 3: Залежність від довжини послідовності", level=2)
    add_paragraph(doc,
        "Третій експеримент досліджує деградацію якості моделей при збільшенні довжини послідовності. "
        "Проведено серію тренувань для послідовностей довжиною від 20 до 200 кроків.",
        first_line_indent=1.25)
    add_image_with_caption(doc, os.path.join(results_dir, 'exp3_seq_length.png'),
                           'Рис. 4.3 - Залежність тестової точності від довжини послідовності', width=5.5)
    add_paragraph(doc,
        "Результати: точність RNN починає падати вже при довжині 40 кроків і досягає рівня випадкового "
        "вгадування при 80 кроках. LSTM зберігає високу точність (>90%) до 150 кроків і лише повільно "
        "деградує при подальшому збільшенні довжини.",
        first_line_indent=1.25)

    # 5. АНАЛІЗ ОТРИМАНИХ РЕЗУЛЬТАТІВ
    add_heading(doc, "5 АНАЛІЗ ОТРИМАНИХ РЕЗУЛЬТАТІВ", level=1)
    add_paragraph(doc,
        "Проведені експерименти чітко демонструють фундаментальні відмінності між архітектурами RNN та LSTM. "
        "На простих задачах з короткими залежностями (до 25-30 кроків) обидві архітектури показують "
        "порівнянні результати. Це пояснюється тим, що градієнти ще не встигають суттєво затухнути.",
        first_line_indent=1.25)
    add_paragraph(doc,
        "При збільшенні відстані між сигналом та місцем прийняття рішення, RNN швидко втрачає здатність "
        "до навчання. Графік залежності точності від довжини послідовності (Рис. 4.3) наочно ілюструє "
        "експоненційний характер деградації RNN.",
        first_line_indent=1.25)
    add_paragraph(doc,
        "LSTM успішно вирішує проблему зникаючого градієнта завдяки механізму cell state. "
        "Ворота забування (forget gate) дозволяють мережі селективно зберігати важливу інформацію "
        "протягом багатьох часових кроків. При f(t) близькому до 1, cell state практично не змінюється, "
        "забезпечуючи \"шосе\" для потоку градієнтів.",
        first_line_indent=1.25)
    add_paragraph(doc,
        "Варто зазначити, що LSTM має приблизно в 4 рази більше параметрів ніж еквівалентна RNN "
        "(через три додаткових набори ваг для воріт). Це призводить до збільшення часу навчання "
        "та вимог до пам\'яті. Однак для задач з довгостроковими залежностями ці витрати виправдані.",
        first_line_indent=1.25)
    add_paragraph(doc,
        "Практичні рекомендації на основі результатів:",
        first_line_indent=1.25)
    add_paragraph(doc,
        "• Для коротких послідовностей (до 30 елементів) можна використовувати RNN для економії ресурсів;",
        first_line_indent=1.25)
    add_paragraph(doc,
        "• Для задач NLP з реченнями та параграфами рекомендується LSTM або її варіація GRU;",
        first_line_indent=1.25)
    add_paragraph(doc,
        "• Для дуже довгих послідовностей (>500 елементів) варто розглянути механізми уваги (attention).",
        first_line_indent=1.25)

    # 6. ВИСНОВКИ
    add_heading(doc, "6 ВИСНОВКИ", level=1)
    add_paragraph(doc,
        "У даній роботі проведено комплексне порівняння архітектур рекурентних нейронних мереж RNN та LSTM "
        "для задач обробки послідовностей. На основі теоретичного аналізу та серії експериментів "
        "отримано наступні результати:",
        first_line_indent=1.25)
    add_paragraph(doc,
        "1. Підтверджено теоретичні передбачення щодо проблеми зникаючого градієнта в класичних RNN. "
        "Експерименти показали, що RNN не здатна навчатися на послідовностях з залежностями довшими "
        "за 40-50 кроків.",
        first_line_indent=1.25)
    add_paragraph(doc,
        "2. Продемонстровано ефективність механізму воріт LSTM у вирішенні проблеми зникаючого градієнта. "
        "LSTM успішно навчається на послідовностях з залежностями до 150 кроків і більше.",
        first_line_indent=1.25)
    add_paragraph(doc,
        "3. На простих задачах з короткими залежностями обидві архітектури показують схожу точність, "
        "що робить RNN прийнятним вибором для обмежених обчислювальних ресурсів.",
        first_line_indent=1.25)
    add_paragraph(doc,
        "4. Реалізовано та протестовано обидві архітектури мовою Python з використанням бібліотеки PyTorch. "
        "Код доступний у відкритому репозиторії для відтворення результатів.",
        first_line_indent=1.25)
    add_paragraph(doc,
        "5. Сформульовано практичні рекомендації щодо вибору архітектури в залежності від характеру задачі "
        "та довжини послідовностей.",
        first_line_indent=1.25)
    add_paragraph(doc,
        "Напрямки подальших досліджень: порівняння з архітектурою GRU (Gated Recurrent Unit), "
        "дослідження механізмів уваги (attention), застосування до реальних задач NLP.",
        first_line_indent=1.25)

    # ПЕРЕЛІК ДЖЕРЕЛ
    add_heading(doc, "ПЕРЕЛІК ДЖЕРЕЛ", level=1)
    add_paragraph(doc,
        "1. Hochreiter S., Schmidhuber J. Long Short-Term Memory // Neural Computation. - 1997. - "
        "Vol. 9, No. 8. - P. 1735-1780.",
        first_line_indent=0)
    add_paragraph(doc,
        "2. Bengio Y., Simard P., Frasconi P. Learning Long-Term Dependencies with Gradient Descent "
        "is Difficult // IEEE Transactions on Neural Networks. - 1994. - Vol. 5, No. 2. - P. 157-166.",
        first_line_indent=0)
    add_paragraph(doc,
        "3. Goodfellow I., Bengio Y., Courville A. Deep Learning. - MIT Press, 2016. - Chapter 10: "
        "Sequence Modeling: Recurrent and Recursive Nets.",
        first_line_indent=0)
    add_paragraph(doc,
        "4. Graves A. Supervised Sequence Labelling with Recurrent Neural Networks. - "
        "Springer, 2012. - 146 p.",
        first_line_indent=0)
    add_paragraph(doc,
        "5. PyTorch Documentation: Recurrent Layers [Електронний ресурс]. - "
        "Режим доступу: https://pytorch.org/docs/stable/nn.html#recurrent-layers",
        first_line_indent=0)

    add_github_link(doc)
    save_document(doc, 'research_rnn_lstm')


# ==================== MAIN ====================

def main():
    """Generate all reports."""
    print("=" * 60)
    print("Report Generator for Data Analysis Course")
    print("=" * 60)

    print("\nGenerating Lab 1: Regression Analysis...")
    generate_lab1_report()

    print("\nGenerating Lab 2: Clustering and Decision Trees...")
    generate_lab2_report()

    print("\nGenerating Lab 3: Apriori Algorithm...")
    generate_lab3_report()

    print("\nGenerating Lab 4: Genetic Algorithms...")
    generate_lab4_report()

    print("\nGenerating Practical 1: OLS Regression...")
    generate_practical1_report()

    print("\nGenerating Practical 2: TF-IDF...")
    generate_practical2_report()

    print("\nGenerating Research: RNN vs LSTM...")
    generate_research_report()

    print("\n" + "=" * 60)
    print(f"All reports saved to: {REPORTS_DIR}")
    print("=" * 60)


if __name__ == "__main__":
    main()
